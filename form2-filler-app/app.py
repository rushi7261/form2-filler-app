import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime
from io import BytesIO
import re

st.set_page_config(page_title="Form 2 Auto‑Filler", layout="centered")
st.title("📄 Form 2 Auto‑Filler – NEXUS")

uploaded_excel = st.file_uploader("Upload Excel with project data", type=["xlsx"])
if uploaded_excel:
    df = pd.read_excel(uploaded_excel)
    st.subheader("📋 Excel Preview")
    st.dataframe(df.head())

    selected_row = st.selectbox("Select row index", df.index)

    if st.button("📝 Generate Word"):
        data = df.loc[selected_row]
        # Ensure all expected keys exist
        replacements = {
            "<Project Name>": str(data.get("Project Name", "")),
            "<Registration Number from View Certificate>": str(data.get("Registration Number", "")),
            "<Promoter Name>": str(data.get("Promoter Name", "")),
            "<Planning Authority Name>": str(data.get("Planning Authority Name", "")),
            "<Date of Certificate>": str(data.get("Date of Certificate", datetime.today().strftime("%Y-%m-%d"))),
            "<Date of Registration>": str(data.get("Date of Registration", datetime.today().strftime("%Y-%m-%d"))),
        }

        try:
            doc = Document("templates/Form 2 (Basic) - NEXUS.docx")
        except Exception as e:
            st.error(f"Template not found: {e}")
            st.stop()

        # Replace in paragraphs
        for p in doc.paragraphs:
            for k, v in replacements.items():
                if k in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(k, v)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for k, v in replacements.items():
                        if k in cell.text:
                            for run in cell.paragraphs[0].runs:
                                run.text = run.text.replace(k, v)

        proj = str(data.get("Project Name", "Project")).strip()
        proj_clean = re.sub(r'[^\w\s-]', '', proj).replace(" ", "_")
        cert_date = pd.to_datetime(data.get("Date of Certificate", datetime.today()))
        quarter_str = cert_date.strftime("%B %Y")
        filename = f"Form 2 - {proj_clean} as on {quarter_str}.docx"

        # Save to buffer instead of disk
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Document ready!")
        st.download_button(f"⬇ Download {filename}", buffer, file_name=filename)
